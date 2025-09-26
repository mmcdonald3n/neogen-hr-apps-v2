from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from pathlib import Path
from .branding import BRAND
def add_page_number(doc):
    section = doc.sections[0]; footer = section.footer
    p = footer.paragraphs[0]; p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    r = run._r; r.append(fldChar1); r.append(instrText); r.append(fldChar2)
def export_docx(md_text: str, title: str, out_path: Path) -> Path:
    doc = Document(); section = doc.sections[0]
    section.top_margin = Inches(0.7); section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8); section.right_margin = Inches(0.8)
    header = section.header
    table = header.add_table(rows=1, cols=2); table.autofit = True
    cell_logo, cell_title = table.rows[0].cells
    try: cell_logo.paragraphs[0].add_run().add_picture(BRAND.logo_path, width=Inches(1.1))
    except Exception: pass
    p = cell_title.paragraphs[0]; p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    r = p.add_run(title); r.font.size = Pt(16); r.bold = True
    add_page_number(doc)
    for block in md_text.split('\n'):
        if block.strip().startswith('#'):
            text = block.lstrip('#').strip()
            p = doc.add_paragraph(); run = p.add_run(text); run.bold = True; run.font.size = Pt(14)
        elif block.strip().startswith('- '):
            doc.add_paragraph(block.strip()[2:], style='List Bullet')
        else:
            if block.strip() == "": doc.add_paragraph("")
            else:
                p = doc.add_paragraph(block.strip()); p.paragraph_format.space_after = Pt(6)
    out_path.parent.mkdir(parents=True, exist_ok=True); doc.save(out_path); return out_path

