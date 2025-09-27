from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from .house_style import BRAND

def markdown_to_plain(md: str) -> str:
    return md.replace("**", "").replace("__", "").replace("\t", "    ")

def export_docx(title: str, md_text: str) -> BytesIO:
    doc = Document()
    try:
        doc.add_picture(BRAND.logo_path, width=Inches(1.4))
    except Exception:
        pass
    h = doc.add_heading(title, level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for line in markdown_to_plain(md_text).splitlines():
        doc.add_paragraph(line)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_markdown(md_text: str) -> BytesIO:
    bio = BytesIO()
    bio.write(md_text.encode('utf-8'))
    bio.seek(0)
    return bio








